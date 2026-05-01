from pathlib import Path
import subprocess
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\mayan\Flutter Projects\smart_portfolio_tracker\mobileclean_app_UI_Documentation_Mar10_Apr24_2026.docx")
REPO = Path(r"C:\Users\mayan\Flutter Projects\mobileclean_app")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = width


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.style = "Table Grid"
    if widths:
        tbl = table._tbl
        tbl_grid = tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            tbl.insert(0, tbl_grid)
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(width.twips)))
            tbl_grid.append(grid_col)
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if widths and idx < len(widths):
                set_cell_width(cell, widths[idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
            cell.margin_top = 80
            cell.margin_bottom = 80
            cell.margin_left = 100
            cell.margin_right = 100
            if row_idx == 0:
                shade_cell(cell, "1F4E5A")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)


def add_table(doc, headers, rows, widths=None):
    key_w = 18
    val_w = 82
    border = "+" + "-" * (key_w + 2) + "+" + "-" * (val_w + 2) + "+"

    def add_mono_run(paragraph, text, bold=False, highlight=False, color=None):
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(7.2)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return run

    for row in rows:
        rendered_lines = [("border", "", "", "")]
        for i, header in enumerate(headers):
            value = str(row[i] if i < len(row) else "")
            value = value.replace("; ", ";\n")
            key_lines = textwrap.wrap(header, key_w) or [""]
            val_lines = []
            for part in value.splitlines():
                val_lines.extend(textwrap.wrap(part, val_w, break_long_words=True, break_on_hyphens=False) or [""])
            height = max(len(key_lines), len(val_lines))
            for idx in range(height):
                k = key_lines[idx] if idx < len(key_lines) else ""
                v = val_lines[idx] if idx < len(val_lines) else ""
                rendered_lines.append(("row", header, k, v))
            rendered_lines.append(("border", "", "", ""))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        for line_idx, (kind, header, key_text, val_text) in enumerate(rendered_lines):
            if line_idx:
                add_mono_run(p, "\n")
            if kind == "border":
                add_mono_run(p, border, color="6B7280")
            else:
                highlight_value = "path" in header.lower() or header.lower() == "changed files"
                add_mono_run(p, "| ", color="6B7280")
                add_mono_run(p, f"{key_text:<{key_w}}", bold=True)
                add_mono_run(p, " | ", color="6B7280")
                add_mono_run(p, f"{val_text:<{val_w}}", highlight=highlight_value)
                add_mono_run(p, " |", color="6B7280")
    return

    # Real Word table implementation kept below for future use. The current
    # artifact renderer collapses Word table cells with long Flutter paths.
    for row in rows:
        table = doc.add_table(rows=len(headers), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "9500")
        tbl_w.set(qn("w:type"), "dxa")
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)
        label_width = Inches(1.45)
        value_width = Inches(5.15)
        tbl_grid = table._tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, tbl_grid)
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for width in (label_width, value_width):
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(width.twips)))
            tbl_grid.append(grid_col)
        for i, header in enumerate(headers):
            label_cell = table.rows[i].cells[0]
            value_cell = table.rows[i].cells[1]
            set_cell_width(label_cell, label_width)
            set_cell_width(value_cell, value_width)
            set_cell_text(label_cell, header, bold=True, color="FFFFFF", size=8.5)
            shade_cell(label_cell, "1F4E5A")
            value = row[i] if i < len(row) else ""
            value = str(value).replace("; ", ";\n")
            set_cell_text(value_cell, value, size=8.5)
            for cell in (label_cell, value_cell):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.05
                    for run in p.runs:
                        run.font.name = "Calibri"
        doc.add_paragraph()


def classify_file(path):
    mapping = [
        ("lib/screens/D1MM3", "D1MM3 Membership Plan"),
        ("lib/local_components/widgets/D1MM3", "D1MM3 Membership Plan"),
        ("lib/local_components/controllers/D1MM3", "D1MM3 Membership Plan"),
        ("lib/api/D1MM3", "D1MM3 Membership Plan API"),
        ("lib/screens/D1MM4", "Dashboard / D1MM4"),
        ("lib/local_components/widgets/D1MM4", "Dashboard / D1MM4"),
        ("lib/local_components/controllers/D1MM4", "Dashboard / D1MM4"),
        ("lib/screens/D1MM10", "Dashboard V2 / Members"),
        ("lib/local_components/widgets/D1MM10", "Dashboard V2 / Members"),
        ("lib/local_components/controllers/D1MM10", "Dashboard V2 / Members"),
        ("lib/api/D1MM10", "Dashboard V2 API"),
        ("lib/screens/D1CC6", "Attendance Verification"),
        ("lib/local_components/widgets/D1CC6", "Attendance Verification"),
        ("lib/local_components/controllers/D1CC6", "Attendance Verification"),
        ("lib/api/D1CC6", "Attendance Verification API"),
        ("lib/screens/D1MM5", "My Profile / Finance Settings"),
        ("lib/local_components/controllers/D1MM5", "My Profile / Finance Settings"),
        ("lib/local_components/models/D1MM5", "My Profile / Finance Settings"),
        ("lib/screens/D1CC7", "Notifications and Reminders"),
        ("lib/local_components/controllers/D1CC7", "Notifications and Reminders"),
        ("lib/api/D1CC7", "Notifications API"),
        ("lib/global_components", "Global Components"),
        ("lib/routes", "Routing"),
        ("lib/utils", "Utilities / Constants"),
        ("lib/config", "Configuration"),
        ("lib/api/D1MM11", "Payment API"),
        ("lib/api/D1MM9", "KYC API"),
    ]
    normalized = path.replace("\\", "/")
    for prefix, label in mapping:
        if normalized.startswith(prefix):
            return label
    return "Other"


def collect_commit_details():
    cmd = [
        "git",
        "-c",
        f"safe.directory={REPO.as_posix()}",
        "log",
        "--since=2026-03-10",
        "--until=2026-04-25 23:59:59",
        "--reverse",
        "--pretty=format:COMMIT%x09%h%x09%ad%x09%an%x09%s",
        "--date=short",
        "--numstat",
        "--",
        "lib",
    ]
    raw = subprocess.check_output(cmd, cwd=str(REPO), text=True, encoding="utf-8", errors="replace")
    commits = []
    current = None
    for line in raw.splitlines():
        if line.startswith("COMMIT\t"):
            parts = line.split("\t", 4)
            current = {
                "hash": parts[1],
                "date": parts[2],
                "author": parts[3],
                "message": parts[4],
                "files": [],
                "insertions": 0,
                "deletions": 0,
                "areas": set(),
            }
            commits.append(current)
            continue
        if not line.strip() or current is None:
            continue
        parts = line.split("\t")
        if len(parts) < 3:
            continue
        add, delete, path = parts[0], parts[1], parts[2]
        current["files"].append(path)
        current["areas"].add(classify_file(path))
        if add.isdigit():
            current["insertions"] += int(add)
        if delete.isdigit():
            current["deletions"] += int(delete)
    return commits


def add_commit_details(doc):
    add_h(doc, "All Commit-wise Changes", 1)
    intro = doc.add_paragraph()
    intro.add_run(
        "This section includes every commit in the supplied timeline, with the actual files changed under lib."
    ).font.size = Pt(10)
    for idx, commit in enumerate(collect_commit_details(), start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{idx}. {commit['date']} - {commit['message']}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(31, 78, 90)
        files_text = "; ".join(commit["files"])
        if len(files_text) > 1300:
            files_text = files_text[:1297] + "..."
        add_table(
            doc,
            ["Commit", "Change Size", "Affected Areas", "Changed Files"],
            [
                (
                    f"{commit['hash']} by {commit['author']}",
                    f"{len(commit['files'])} files changed, {commit['insertions']} insertions, {commit['deletions']} deletions",
                    ", ".join(sorted(commit["areas"])),
                    files_text,
                )
            ],
        )


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(31, 78, 90)
    return p


def add_meta(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " : ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(value).font.size = Pt(10)


def add_component_section(doc, global_rows, local_rows):
    add_h(doc, "Global Components Used", 3)
    add_table(
        doc,
        ["Component Name", "File Path", "Description"],
        global_rows,
        [Inches(1.65), Inches(2.65), Inches(2.2)],
    )
    add_h(doc, "Local Components Used", 3)
    add_table(
        doc,
        ["Component Name", "File Path", "Description"],
        local_rows,
        [Inches(1.65), Inches(2.65), Inches(2.2)],
    )


def add_page_section(doc, number, title, path, desc, global_rows, local_rows, notes=None):
    add_h(doc, f"{number}. {title}", 2)
    add_table(doc, ["Field", "Details"], [("Path", path), ("Description", desc)])
    if notes:
        add_h(doc, "Code Analysis Notes", 3)
        add_table(doc, ["Note", "Details"], [(f"Note {i}", item) for i, item in enumerate(notes, 1)])
    add_component_section(doc, global_rows, local_rows)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("UI Documentation")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 90)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("mobileclean_app | D1MM3, Dashboard V2, Attendance Verification, My Profile, Notifications")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Details"],
        [
            ("Project Path", r"C:\Users\mayan\Flutter Projects\mobileclean_app"),
            ("Documentation Scope", "Commits from Mar 10, 2026 to Apr 24, 2026 by mayank74-tech"),
            ("Framework", "Flutter with GetX, Dio/HTTP APIs, SharedPreferences/Secure Storage, and reusable global/local widgets"),
        ],
    )
    doc.add_paragraph()

    add_h(doc, "Commit Summary", 1)
    add_table(
        doc,
        ["Date", "Commit Focus", "Implementation Area"],
        [
            ("Mar 10-11, 2026", "Global component and screen component centralization", "Reusable widgets, navigation/header, shared UI structure"),
            ("Mar 12-18, 2026", "Dashboard V2 work", "Dashboard layout, profile completion cards, accessibility shortcuts"),
            ("Mar 19-20, 2026", "UI bug fixes", "Spacing, layout behavior, visual state corrections"),
            ("Mar 24-28, 2026", "D1MM3 Membership Plan module", "Plans, plan selection, membership creation, KYC/payout flow"),
            ("Mar 30-Apr 8, 2026", "Attendance verification", "QR/PIN sharing, scan and connect, live check-in list"),
            ("Apr 9-11, 2026", "Dashboard V2 completion", "Dashboard info containers, member tiles, all members page"),
            ("Apr 13-18, 2026", "My Profile page", "Profile overview, personal details, bank account UI"),
            ("Apr 21, 2026", "Notification and Reminder page", "Notification bottom sheet, controller/service, transaction/general settings UI"),
            ("Apr 23, 2026", "KYC order and membership plan bug fixes", "Payment order flow, membership plan model/API fixes"),
            ("Apr 24, 2026", "UI bugs", "Notification UI, profile pages, payment popup, header state"),
        ],
        [Inches(1.35), Inches(2.25), Inches(3.0)],
    )

    add_commit_details(doc)

    add_h(doc, "For Global Components, refer lib/global_components", 1)

    common_globals = [
        ("Header", "lib/global_components/widgets/navigation/header.dart", "Reusable module header with back button, module name, profile and notification affordances."),
        ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Shared bottom navigation used across dashboard, profile, attendance and finance screens."),
        ("PrimaryButton", "lib/global_components/widgets/primary_button.dart", "Reusable primary CTA button with optional icon support."),
        ("SecondaryButton", "lib/global_components/widgets/secondary_button.dart", "Reusable secondary CTA button for alternative actions."),
        ("Color/Image/Text Constants", "lib/utils/color_constants.dart, image_constants.dart, centeral_text_style .dart", "Centralized styling, assets, and typography values used to keep screens consistent."),
    ]

    add_page_section(
        doc,
        1,
        "D1MM3 - Subscription Plans Page",
        "lib/screens/D1MM3/subscription_plans_page.dart",
        "Plan selection screen where the manager chooses monthly or quarterly membership options, reviews plan cards, accepts plan-switch consent, and opens payment/contact flows.",
        [
            common_globals[0],
            ("TabNav", "lib/global_components/widgets/navigation/tab_nav.dart", "Global tab component used for Monthly and Quarterly plan switching."),
            ("PrimaryButton", "lib/global_components/widgets/primary_button.dart", "CTA for plan purchase or Contact Us flow."),
            ("PaymentPopup", "lib/global_components/widgets/popup/payment_popup.dart", "Reusable payment modal used after plan selection."),
        ],
        [
            ("PricingPlansSection", "lib/local_components/widgets/D1MM3/pricing_plans_section.dart", "Renders selectable membership plan cards and plan-specific pricing states."),
            ("PlanChart", "lib/local_components/widgets/D1MM3/plan_chart.dart", "Visual support component for plan comparison and selected plan details."),
            ("SubscriptionPlanController", "lib/local_components/controllers/D1MM3/subscription_plan_controller.dart", "Holds selected plan state, API data, and enterprise/contact sheet behavior."),
            ("TickBoxController", "lib/local_components/controllers/D1MM4/tick_box_controller.dart", "Controls optional switch-to-individual-plan consent checkbox."),
        ],
        [
            "Uses GetX observables for selected plan state and conditional CTA rendering.",
            "Apr 23 membership plan bug work touched API/model/payment flow files linked to this screen.",
        ],
    )

    add_page_section(
        doc,
        2,
        "D1MM3 - Plan Creation Form Page",
        "lib/screens/D1MM3/plan_creation_form_page.dart",
        "Membership creation form where the manager enters membership name, type, objectives, session/period details, pricing and other plan configuration before preview.",
        [
            common_globals[0],
            ("TextInputFiled", "lib/global_components/widgets/D1MM2/text_form_filed.dart", "Reusable text input used for plan name, sessions, validity and pricing fields."),
            ("CustomDropDown", "lib/global_components/widgets/D1MM2/custom_drop_down.dart", "Dropdown control reused from D1MM2 form components."),
            ("BubbleTab", "lib/global_components/widgets/D1MM2/bubble_tab.dart", "Chip-like multi-option selector used for membership objectives."),
            ("RadioButtonTab", "lib/global_components/widgets/D1MM2/radio_button_tab.dart", "Segmented selection for period-based versus session-based plan."),
            ("CustomCheckBox", "lib/global_components/widgets/D1MM2/custom_check_box.dart", "Checkbox control used in plan setup/confirmation areas."),
        ],
        [
            ("PlanCreationFormController", "lib/local_components/controllers/D1MM3/plan_creation_form_controller.dart", "Owns text controllers, selected membership type, objectives and navigation actions."),
            ("MultiSelectDropdown", "lib/local_components/widgets/D1MM3/multi_select_dropdown.dart", "D1MM3 local multi-select widget for configurable form selections."),
            ("ProgressBar", "lib/local_components/widgets/D1MM3/progress_bar.dart", "Shows the user's current step across the plan creation flow."),
        ],
        [
            "Conditional form sections are driven by the selected membership type.",
            "The screen uses a blurred background card layout matching the dashboard visual language.",
        ],
    )

    add_page_section(
        doc,
        3,
        "D1MM3 - Plan Overview and Payout Flow",
        "lib/screens/D1MM3/plan_form_overview_page.dart; lib/screens/D1MM3/payout_overview_page.dart",
        "Review and payout acceptance screens used after plan creation. They summarize entered plan details, show payout/tax information and require user consent before continuing.",
        [
            common_globals[0],
            ("PrimaryButton", "lib/global_components/widgets/primary_button.dart", "Used for accept/continue actions."),
            ("CustomCheckBox", "lib/global_components/widgets/D1MM2/custom_check_box.dart", "Used for tax consent and terms acceptance states."),
        ],
        [
            ("PlanFormOverviewController", "lib/local_components/controllers/D1MM3/plan_form_overview_controller.dart", "Controls preview/overview actions."),
            ("PayoutOverviewController", "lib/local_components/controllers/D1MM3/payout_overview_controller.dart", "Controls back/accept actions and consent states."),
            ("Overview widgets", "lib/local_components/widgets/D1MM3/overview/*.dart", "Reusable overview cards, rows, section titles, chips, dividers and consent checkbox."),
            ("ProgressBar", "lib/local_components/widgets/D1MM3/progress_bar.dart", "Keeps the multi-step flow visible to the user."),
        ],
        [
            "Apr 6 work adjusted Obx usage and direct observable reads to reduce reactive UI errors.",
            "The flow intentionally blocks back navigation with WillPopScope in the creation sequence.",
        ],
    )

    add_page_section(
        doc,
        4,
        "D1MM3 - KYC Verification and Plan Listing",
        "lib/screens/D1MM3/kyc_verification_page.dart; lib/screens/D1MM3/plan_listing.dart",
        "KYC verification gate and membership listing screens connected to plan purchase/management. The KYC order bug fixed on Apr 23 relates to payment and membership APIs used around this flow.",
        [
            common_globals[0],
            ("PrimaryButton", "lib/global_components/widgets/primary_button.dart", "Primary action control used for verification and navigation."),
            ("PaymentStatusPopup", "lib/global_components/widgets/popup/payment_status_popup.dart", "Payment feedback component used by related plan flows."),
        ],
        [
            ("ProgressBarController", "lib/local_components/controllers/D1MM3/progress_bar_controller.dart", "Tracks flow step progress around KYC and payout steps."),
            ("CardController", "lib/local_components/controllers/D1MM4/custom_card_controller.dart", "Provides current onboarding/verification step state used by KYC and dashboard."),
            ("MembershipPlanListing", "lib/local_components/widgets/D1MM3/membership_plan_listing.dart", "Renders saved membership plan items."),
            ("MembershipPlansTile", "lib/local_components/widgets/D1MM3/membership_plans_tile.dart", "Individual membership plan row/card UI."),
        ],
    )

    add_page_section(
        doc,
        5,
        "Dashboard V2 - Manager Home Page",
        "lib/screens/D1MM4/v1_dashboard.dart",
        "Dashboard landing page updated during Dashboard V2 work. It shows profile/onboarding completion cards, access shortcuts and switches to richer dashboard information once setup progress reaches the required step.",
        [
            ("Color/Image Constants", "lib/utils/color_constants.dart; lib/utils/image_constants.dart", "Shared background, overlay and color palette."),
            common_globals[4],
        ],
        [
            ("Header", "lib/local_components/widgets/D1MM4/header.dart", "Dashboard-specific header with profile/notification state."),
            ("CustomCard", "lib/local_components/widgets/D1MM4/custom_card.dart", "Profile/setup completion card shown at the top of the dashboard."),
            ("Component2 Accessibility Box", "lib/local_components/widgets/D1MM4/accessibility_box.dart", "Shortcut grid for bank account, plan, and member workflows based on enabled state."),
            ("Component3 QuickInfo", "lib/local_components/widgets/D1MM4/quickinfo.dart", "Pre-completion quick info block."),
            ("Component4 HorizontalInfoBox", "lib/local_components/widgets/D1MM4/horizontal_info_box.dart", "Horizontal dashboard metrics block."),
            ("DashboardInfoContainer", "lib/local_components/widgets/D1MM10/dashboard_info_container.dart", "Dashboard V2 information container shown after setup progress."),
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation shared across main modules."),
        ],
        [
            "Dashboard V2 work moved from the earlier v2_dashboard.dart into v1_dashboard.dart and D1MM10 widgets.",
            "CardController.fetchInitialStep controls which dashboard blocks become available.",
        ],
    )

    add_page_section(
        doc,
        6,
        "Dashboard V2 - All Members Page",
        "lib/screens/D1MM10/all_members_page.dart",
        "Member list page opened from the dashboard shortcut. It supports All, Active and Pending filters with selection behavior for pending activation limits.",
        [
            ("QuickInfo", "lib/local_components/widgets/D1MM10/quick_info_container.dart", "Dashboard quick information widget reused in member views."),
            common_globals[4],
        ],
        [
            ("LiveCheckInController", "lib/local_components/controllers/D1CC6/live_checkin_controller.dart", "Provides filtered member list and tab state."),
            ("HorizontalInfoBoxController", "lib/local_components/controllers/D1MM10/horizontal_info_box_controller.dart", "Manages pending member selection and activation limits."),
            ("MemberTile", "lib/local_components/widgets/shared/member_tile.dart", "Reusable member row/card used for lists and popups."),
            ("QuickInfo", "lib/local_components/widgets/D1MM10/quick_info_container.dart", "Filter-aware summary widget."),
        ],
    )

    add_page_section(
        doc,
        7,
        "D1CC6 - Scan and Connect Page",
        "lib/screens/D1CC6/qr_scanning.dart",
        "Attendance verification page where a manager can display/share a QR code, provide a manual PIN, download/share QR assets and navigate to live check-ins.",
        [
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation for returning to main modules."),
            ("URL Launcher / Share helpers", "pubspec.yaml dependencies: url_launcher, share_plus, qr_flutter, gal", "Packages supporting share, open link and QR related actions."),
        ],
        [
            ("ScanConnectController", "lib/local_components/controllers/D1CC6/attendance_verification_controller.dart", "Loads QR/PIN state, business details, share/download behavior and error states."),
            ("_QrSection", "lib/screens/D1CC6/qr_scanning.dart", "Inline local widget for QR visible, loading, unavailable and blurred states."),
            ("_ShareBottomSheet", "lib/screens/D1CC6/qr_scanning.dart", "Local bottom sheet for sharing QR/PIN through supported channels."),
            ("QrDownloadTemplate", "lib/screens/D1CC6/qr_scanning.dart", "Download/share template rendered from the QR state."),
            ("QrSaveDialog", "lib/screens/widgets/qr_Save_dialog.dart", "Dialog shown while saving QR assets."),
            ("ScanningUnavailable", "lib/screens/widgets/scanning_unavailable.dart", "Fallback UI for unavailable scanning state."),
        ],
        [
            "Mar 30-Apr 8 commits introduced attendance API, QR UI, scan unavailable fallback and live check-in models.",
            "The page uses local inline widgets heavily, so component documentation names those private widgets explicitly.",
        ],
    )

    add_page_section(
        doc,
        8,
        "D1CC6 - Live Check-In Page",
        "lib/screens/D1CC6/live_checkin_screen.dart",
        "Live member check-in list with summary information and filters for all, active only and expiring members.",
        [
            ("Color/Image Constants", "lib/utils/color_constants.dart; lib/utils/image_constants.dart", "Background layering and common colors."),
        ],
        [
            ("LiveCheckInController", "lib/local_components/controllers/D1CC6/live_checkin_controller.dart", "Stores member list, filter type and popup tab state."),
            ("QuickInfo", "lib/local_components/widgets/D1MM10/quick_info_container.dart", "Displays summary count information based on selected filter."),
            ("MemberTile", "lib/local_components/widgets/shared/member_tile.dart", "Renders member entries and supports popup/detail behavior."),
            ("StatusBadge", "lib/local_components/widgets/D1CC6/status_badge.dart", "Reusable status chip for active, pending, expiring, expired and renewal states."),
        ],
    )

    add_page_section(
        doc,
        9,
        "D1MM5 - My Profile Overview Page",
        "lib/screens/D1MM5/profile_screen.dart",
        "My Profile landing page showing manager profile data, verification badge, account settings sections and logout action.",
        [
            common_globals[0],
            ("ProfileController", "lib/global_components/controllers/profile_controller.dart", "Fetches and exposes profile data used by the profile card."),
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation across manager modules."),
        ],
        [
            ("_ProfileCard", "lib/screens/D1MM5/profile_screen.dart", "Inline card showing avatar, name, role and verification status."),
            ("_SettingsSectionsCard", "lib/screens/D1MM5/profile_screen.dart", "Inline settings panel with links to profile sub-pages."),
            ("_SettingsItem", "lib/screens/D1MM5/profile_screen.dart", "Reusable inline list item for Personal Details, Bank Account, Transactions and General Settings."),
            ("_LogoutButton", "lib/screens/D1MM5/profile_screen.dart", "Inline logout action."),
            ("VerificationStatus", "lib/local_components/widgets/D1MM4/verification_status.dart", "Verification badge reused from dashboard components."),
        ],
        [
            "Apr 13-Apr 18 commits added profile overview and linked sub-pages.",
            "Apr 24 UI bug commit touched profile, personal detail, bank account and transaction screens.",
        ],
    )

    add_page_section(
        doc,
        10,
        "D1MM5 - Personal Details Update Page",
        "lib/screens/D1MM5/personal_detail_screen.dart",
        "Personal information update screen with profile image picker, editable fields, read-only fields, gender/date/dropdown/language controls and loading state.",
        [
            common_globals[0],
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation."),
            ("Image Picker / Dotted Border", "pubspec.yaml dependencies: image_picker, dotted_border", "Packages supporting profile image upload UI."),
        ],
        [
            ("PersonalDetailsController", "lib/local_components/controllers/D1MM5/personal_detail_controller.dart", "Owns form controllers, loading state, gender/date/language values and image selection."),
            ("_FormCard", "lib/screens/D1MM5/personal_detail_screen.dart", "Local form container for personal details."),
            ("_ProfileImagePicker", "lib/screens/D1MM5/personal_detail_screen.dart", "Local image selection and preview UI."),
            ("_EditableField / _ReadOnlyField / _PhoneField", "lib/screens/D1MM5/personal_detail_screen.dart", "Inline reusable field widgets for form rows."),
            ("_GenderSelector / _DateField / _DropdownField / _LanguageSelector", "lib/screens/D1MM5/personal_detail_screen.dart", "Local controls for structured personal detail inputs."),
            ("UpdationPopup", "lib/local_components/widgets/popup/updation_popup.dart", "Popup used to confirm or show update feedback in the profile flow."),
        ],
    )

    add_page_section(
        doc,
        11,
        "D1MM5 - Bank Account Detail Page",
        "lib/screens/D1MM5/bank_account_detail_screen.dart",
        "Bank account screen showing credited amount, transaction summary bars, linked bank account cards, visibility toggle and detail bottom sheet.",
        [
            common_globals[0],
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation."),
        ],
        [
            ("_BankAccount", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Static local model currently backing the bank account UI."),
            ("_TabBarWidget", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Local Today/This Week/This Month filter tabs."),
            ("_AmountSection", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Credited payment amount area with visibility state."),
            ("_TransactionCard", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Summary transaction chart/card."),
            ("_BankAccountCard", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Linked bank account card."),
            ("_BankDetailSheet", "lib/screens/D1MM5/bank_account_detail_screen.dart", "Bottom sheet with detailed bank account information."),
        ],
        ["The reference-style note applies here: live bank APIs appear pending; the current screen uses static data placeholders."],
    )

    add_page_section(
        doc,
        12,
        "D1MM5 - Transactions Page",
        "lib/screens/D1MM5/transaction_screen.dart",
        "Transaction history screen with All, Receipts and Payments filters, mock transaction data, status styling and transaction details dialog.",
        [
            common_globals[0],
            ("CustomBottomNavigationBar", "lib/local_components/widgets/D1MM4/custom_navigation_bar.dart", "Bottom navigation."),
        ],
        [
            ("BankAccountModel", "lib/local_components/models/D1MM5/bank_account_model.dart", "Model/enums for transaction category and status."),
            ("_FilterRow / _PillTab", "lib/screens/D1MM5/transaction_screen.dart", "Local filtering UI for transaction categories."),
            ("_TxTile", "lib/screens/D1MM5/transaction_screen.dart", "Local transaction list tile."),
            ("TransactionDetailsDialog", "lib/screens/D1MM5/transaction_screen.dart", "Dialog that displays expanded transaction details."),
        ],
        ["The screen is structured for future controller/API replacement; current data source is _mockTransactions."],
    )

    add_page_section(
        doc,
        13,
        "D1MM5 - General Settings Page",
        "lib/screens/D1MM5/general_setting_screen.dart",
        "General settings screen for promotional notification preference and account deactivation flow with confirmation/success dialogs.",
        [
            common_globals[0],
            ("AuthController", "lib/global_components/controllers/auth_controller.dart", "Used for logout from the deactivation success dialog."),
        ],
        [
            ("_CustomDialog", "lib/screens/D1MM5/general_setting_screen.dart", "Local reusable dialog for notification toggle and deactivation confirmation."),
            ("Notification toggle state", "lib/screens/D1MM5/general_setting_screen.dart", "Local isNotificationOn state controls promotional notification preference."),
            ("Deactivation flow", "lib/screens/D1MM5/general_setting_screen.dart", "Local dialog sequence tracks pending state and follow-up user actions."),
        ],
    )

    add_page_section(
        doc,
        14,
        "D1CC7 - Notification and Reminder Page",
        "lib/screens/D1CC7/notification_screen.dart",
        "Notification bottom sheet with categories, loading/error/empty states, pull-to-refresh and mark-as-read behavior.",
        [
            ("Color Constants", "lib/utils/color_constants.dart", "Shared dark sheet and accent colors."),
            ("SharedPreferences", "pubspec.yaml dependency: shared_preferences", "Used by notification/auth related state."),
        ],
        [
            ("NotificationController", "lib/local_components/controllers/D1CC7/notification_controller.dart", "Fetches notifications, exposes filtered list, handles mark-as-read and loading/error state."),
            ("NotificationService", "lib/api/D1CC7/notification_service.dart", "API service for notification data."),
            ("NotificationModel", "lib/screens/D1CC7/notification_screen.dart", "Local UI model used by the list."),
            ("NotificationCategory", "lib/screens/D1CC7/notification_screen.dart", "Enum that maps notification categories to labels, icons and colors."),
            ("TopDownNotificationWidget", "lib/screens/widgets/D1CC7/top_down_notification_widget.dart", "Top-down notification widget added in the Apr 24 UI bug work."),
        ],
        [
            "Apr 21 introduced notification service/controller; Apr 24 expanded the screen and top-down widget.",
            "The sheet height is 75 percent of the viewport and uses a drag handle plus close button.",
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
